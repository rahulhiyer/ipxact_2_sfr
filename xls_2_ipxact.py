'''
This Program takes in the XLS file which is defined by SET team to IP-XACT format of only the SFR part
'''


import xlrd
import xlsxwriter
import re
import docx
import time
#import openpyxl
import logging
import sys
import os
import shutil
import datetime
import sfr_code

#logging.basicConfig(level=logging.DEBUG,format='%(asctime)s - %(levelname)-10s - %(message)s')


class Register(object):
    total_registers = 0
    register_address = []
    register_name = []
    
    def __init__(self,f_handler,sfr_width=32):
        self.__register    = {"name":"","offset":"","reset":"","description":"","s_array":"","e_array":"","port":"","group_name":"","next_Addr":"","drepeat":""}
        self.__field_names = [False] * 32
        self.__reset_values = [0] * 32
        self.__field       = {"name":"","lsb":"","msb":"","access_tag":"","reset_value":"","sysrdl_precedence":"","hdl_path":"","soft_reset_mask":"","description":"","RESERVED":"","volatile":"","sysrdl_we":"","Access":"","UVM_ACCESS":"","ReadAction":"","modifiedWriteValue":"","volatile":"","sfr_doc_type":""}
        self.__field_array =[]
        self.__unique_field_names = []
        self.__possible_access_tag = ["RWi","RdWi","RW","RdW","RcWi","RsWi","RcW","RsW","RWc","RdWc","RWs","RdWs","RcWs","RsWc","RW0i1c","RdW0i1c","RW0i1s","RdW0i1s","RW0i1t","RdW0i1t","RW0c1i","RdW0c1i","RW0s1i","RdW0s1i","RW0t1i","RdW0t1i","RcW0i1s","RdcW0i1s","RsW0i1c","RcW0s1i","RdcW0s1i","RsW0c1i","RiW","RiWc","RiWs","RWsty","RdWsty","RiWsty"]        
        self.__sfr_width = sfr_width
        self.f_handler = f_handler
        self.__nibbles     = sfr_width/4
        self.hexaPattern = re.compile(r'0x([0-9a-fA-F]+)')
        self.hexaPatternE2E = re.compile(r'^0x([0-9a-fA-F]+)$')

    
    def writing_sfr_doc(self,doc,offset_address=0):
        
        widths= [docx.shared.Cm(2),docx.shared.Cm(2),docx.shared.Cm(2),docx.shared.Cm(2)]        
        p = doc.add_paragraph("Offset "+str(hex(int(re.sub("0x","",self.__register["offset"]),base=16)+offset_address))+"h:"+self.__register["name"]+":"+self.__register["description"]) #, style='ListNumber')
        #p.add_run('bold').bold = True
        table=doc.add_table(rows=len(self.__unique_field_names)+1,cols=4,style="ColorfulList-Accent5")
        table.style = 'TableGrid'
        table.allow_autofit = True
        row=table.rows[0]
        #row.cells[0].width =widths[0]
        row.cells[0].text ="Bit"
        #row.cells[1].width =widths[1]
        row.cells[1].text ="Type"
        #row.cells[2].width =widths[2]
        row.cells[2].text ="Reset"
        #row.cells[3].width =widths[3]
        row.cells[3].text ="Description"
        unique_names_from_msb_lsb = self.__unique_field_names[::-1] #reversing the array. [31:0]
        for i in range(len(unique_names_from_msb_lsb)):
            fields = [d for d in self.__field_array if d["name"] is unique_names_from_msb_lsb[i]]            
            assert len(fields) == 1,"Fields Names should be unique : Register : %s, Field Name : %s" %(self.__register["name"], d["name"])
            field = fields[0]            
            row = table.rows[i+1]
            #row.cells[0].width =widths[0]
            if field["lsb"] == field["msb"]:
                row.cells[0].text ="["+str(field["lsb"])+"]"
            else:
                row.cells[0].text ="["+str(field["msb"])+":"+str(field["lsb"])+"]"
            #row.cells[1].width =widths[1]
            row.cells[1].text =field["sfr_doc_type"]
            #row.cells[2].width =widths[2]
            row.cells[2].text =field["reset_value"] #str(int(field["msb"])-int(field["msb"])+1)+"'h"+field["access_tag"]
            #row.cells[3].width =widths[3]
            row.cells[3].text = field["name"]+": "+field["description"]

        p = doc.add_paragraph("\n\n")
        doc.save("sfr_script.docx")
                          

                                        
    def update_register_name(self, name, offset, array, description, port, group_name, drepeat):
       """
       This Program updates the register fields.
       Following checks are added
       1. A check is made if the name does not have expression
       2. A check is made if the offset is mentioned in 0x Hexadecimal format
       3. A check if the register has description
       """
       m = re.compile(r'(.*?)(\$)?$') #Greedy Search
       naming = m.search(name)
       assert naming, "Some issue in naming"
       self.__register["name"] = naming.group(1)
       if naming.group(2) == "$":
           logging.debug("This is array Reg Name : %s" %(name))
           
           reg_ex_hexa = re.compile(r'0x([0-9a-f]+)\s*,\s*0x([0-9a-f]+)',re.I)
           reg_ex_comma_next_addr = reg_ex_hexa.search(offset)
           assert reg_ex_comma_next_addr, "Offset should be 0x?,0x? format Offset : %s, Register Name : %s" %(offset, self.__register["name"])
           next_addr = int(reg_ex_comma_next_addr.group(2),16)
           
           m = re.match(r'(\d+)\.{2}(\d+)',array)
           assert m, "Arrays should be of the format 0..13 : %s Register Name : %s" %(array, self.__register["name"])
           self.__register["s_array"]     = int(m.group(1))
           self.__register["e_array"]     = int(m.group(2))
           logging.debug("Next Addr is multiple of %s" %(next_addr))
       else:
           assert re.match(self.hexaPatternE2E,offset), "Offset should be 0x? format Offset : %s, Register Name : %s" %(offset, self.__register["name"])
           next_addr = "4"
           self.__register["s_array"]     = ""
           self.__register["e_array"]     = ""
           
       self.__register["next_addr"] = next_addr
        
       m = re.match(self.hexaPattern,offset)
       if(m):
           self.__register["offset"]      = m.group(1)
       else:
           assert False,"Offset should start with 0x and should be in hexadecimal : Register Name : "+name+", Offset : "+offset        
              
       self.__register["description"] = description    
       self.__register["port"]        = port
       self.__register["group_name"]  = group_name
       self.__register["drepeat"]     = drepeat
       




    def print_register_names(self, file):        
        if (self.__register["s_array"] == "" or (self.__register["drepeat"] != "")):
            self.writing_into_xml_file(file, "<spirit:register>")
            self.writing_into_xml_file(file, "      <spirit:name>"+self.__register["name"]+"</spirit:name>")
            self.writing_into_xml_file(file, "      <spirit:addressOffset>0x"+self.__register["offset"]+"</spirit:addressOffset>")
            self.writing_into_xml_file(file, "      <spirit:description>"+self.__register["description"]+"</spirit:description>")
            self.print_field_array(file)
            self.writing_into_xml_file(file, "</spirit:register>")
        else:
            logging.debug("Register name is %s" %(self.__register["name"]))
            for i in range(self.__register["e_array"] - self.__register["s_array"] + 1):
                offset = int(self.__register["offset"],16)
                offset = offset + i*self.__register["next_addr"]
                offset = hex(offset) #hex() will return prefix with 0x
                self.writing_into_xml_file(file, "<spirit:register>")
                self.writing_into_xml_file(file, "      <spirit:name>"+self.__register["name"]+"_"+str(int(self.__register["s_array"]) + i)+"</spirit:name>")
                self.writing_into_xml_file(file, "      <spirit:addressOffset>"+offset+"</spirit:addressOffset>")
                self.writing_into_xml_file(file, "      <spirit:description>"+self.__register["description"]+"-"+str(i)+"</spirit:description>")
                self.print_field_array(file)            
                self.writing_into_xml_file(file, "</spirit:register>")


    def print_field_array(self, file):
        """
        This program prints all the fields from High to Low. i.e, 31 to 0
        """        
        unique_names_from_msb_lsb = self.__unique_field_names[::-1] #reversing the array. [31:0]
        for i in unique_names_from_msb_lsb:            
            fields = [d for d in self.__field_array if d["name"] is i]
            assert len(fields) == 1,"Fields Names should be unique : Register : %s, field : %s" %(self.__register["name"], fields)
            self.writing_into_xml_file(file, "      <spirit:field>")
            self.writing_into_xml_file(file, "            <spirit:name>"+fields[0]["name"]+"</spirit:name>")
            self.writing_into_xml_file(file, "            <spirit:description>"+fields[0]["description"]+"</spirit:description>")
            self.writing_into_xml_file(file, "            <spirit:bitOffset>"+str(fields[0]["lsb"])+"</spirit:bitOffset>")
            self.writing_into_xml_file(file, "            <spirit:bitWidth>"+str(int(fields[0]["msb"])-int(fields[0]["lsb"])+1)+"</spirit:bitWidth>")
            self.writing_into_xml_file(file, "            <spirit:volatile>"+fields[0]["volatile"]+"</spirit:volatile>")
            self.writing_into_xml_file(file, "            <spirit:access>"+fields[0]["Access"]+"</spirit:access>")
            if (fields[0]["modifiedWriteValue"] != '') :
                self.writing_into_xml_file(file, "            <spirit:modifiedWriteValue>"+fields[0]["modifiedWriteValue"]+"</spirit:modifiedWriteValue>")
            if (fields[0]["ReadAction"] != '') :
                self.writing_into_xml_file(file, "            <spirit:ReadAction>"+fields[0]["ReadAction"]+"</spirit:ReadAction>")
            self.print_parameters(file, fields[0], self.__register["name"])
            self.writing_into_xml_file(file, "      </spirit:field>")


    def print_parameters(self, file, fields, reg_name):
        self.writing_into_xml_file(file, "            <spirit:parameters>")
        self.writing_into_xml_file(file, "                  <spirit:parameter>")
        self.writing_into_xml_file(file, "                        <spirit:name>_resetValue_</spirit:name>")
        self.writing_into_xml_file(file, "                        <spirit:value>"+fields["reset_value"]+"</spirit:value>")
        self.writing_into_xml_file(file, "                  </spirit:parameter>")
        if (fields["sysrdl_precedence"] != '') :
            self.writing_into_xml_file(file, "                  <spirit:parameter>")
            self.writing_into_xml_file(file, "                        <spirit:name>_sysrdl_precedence_</spirit:name>")
            self.writing_into_xml_file(file, "                        <spirit:value>"+"sw"+"</spirit:value>")
            self.writing_into_xml_file(file, "                  </spirit:parameter>")


        input_port = ("hw_wr_"+reg_name+"_"+fields["name"]).lower()
        if (fields["modifiedWriteValue"] != '') :    
            self.writing_into_xml_file(file, "                  <spirit:parameter>")
            self.writing_into_xml_file(file, "                        <spirit:name>_sysrdl_we_</spirit:name>")
            self.writing_into_xml_file(file, "                        <spirit:value>"+input_port+"_en"+"</spirit:value>")
            self.writing_into_xml_file(file, "                  </spirit:parameter>")
        #if (fields["volatile"] == 'true') :
            self.writing_into_xml_file(file, "                  <spirit:parameter>")
            self.writing_into_xml_file(file, "                        <spirit:name>_inputPort_</spirit:name>")
            self.writing_into_xml_file(file, "                        <spirit:value>"+input_port+"</spirit:value>")
            self.writing_into_xml_file(file, "                  </spirit:parameter>")

         
        self.writing_into_xml_file(file, "            </spirit:parameters>")


    def writing_into_xml_file(self,file,statement):
        '''
        Writing into xml file
        '''        
        file.write(statement)
        file.write("\n")
        logging.debug(statement)


    def printing(self):        
        logging.info("Register name : %s, Offset Address : %s, Port : %s" %(self.__register["name"], self.__register["offset"], self.__register["port"]))
        
                      
    def update_field_name(self,name,lsb,msb,access_tag,reset_value,sysrdl_precedence,hdl_path,soft_reset_mask,description,reserved):
        if (name in self.__field_names):
            assert "Register Field is repeated"
        offset_reset_format = re.compile(r'^0x[0-9A-Fa-f]{1,'+str(self.__nibbles)+'}')
        assert not re.match(r'\s',name),"Cannot have space in field name"
        assert not offset_reset_format.search(offset),"Offset should start with 0x"
        assert not offset_reset_format.search(reset_value),"reset should start with 0x"
        #assert re.match(r'[0-9]',lsb),"LSB should be decimal value"
        #assert re.match(r'[0-9]',msb),"MSB should be decimal value"
        assert access_tag in self.__possible_access_tag,"In-Correct Access Tag : Register name: %s, Field Name : %s, Access Tag : %s" %(self.__register["name"], name, access_tag)
        assert int(lsb) < self.__sfr_width and int(lsb) >= 0, "LSB value should be valid : Register: %s, Field Name : %s, Access Tag : %s" %(self.__register["name"], name, access_tag)
        assert int(msb) < self.__sfr_width and int(msb) >= 0, "MSB value should be valid : Register: %s, Field Name : %s, Access Tag : %s" %(self.__register["name"], name, access_tag)
        self.__field_array.append(dict(self.__field))
        self.__field_array[-1]["name"] = name
        self.__field_array[-1]["lsb"] = lsb
        self.__field_array[-1]["msb"] = msb
        self.__field_array[-1]["access_tag"] = access_tag
        self.__field_array[-1]["reset_value"] = reset_value
        self.__field_array[-1]["sysrdl_precedence"] = sysrdl_precedence
        self.__field_array[-1]["hdl_path"] = hdl_path
        self.__field_array[-1]["soft_reset_mask"] = soft_reset_mask
        self.__field_array[-1]["description"] = description
        self.__field_array[-1]["RESERVED"] = reserved
        if re.match(r'Rd',self.__field_array[-1]["access_tag"]):
            self.__field_array[-1]["volatile"] = "true"
        elif reserved == "RESERVED":
            self.__field_array[-1]["volatile"] = "true"
        else:
            self.__field_array[-1]["volatile"] = "false"

        '''For Read-Only, except Reserved, there should be IO port'''
        if re.match(r'Rd',access_tag) and re.match(r'RO',access_tag) and reserved == "":
            self.__field_array[-1]["volatile"] = "true"

            
        #self.__field_array[-1]["sysrdl_we"] = re.match(r'd',access_tag) and ((re.match(r'1s',access_tag) or (re.match(r'1c',access_tag) or (re.match(r'0s',access_tag) or (re.match(r'0c',access_tag))
        if (reserved == "RESERVED"):
            self.__field_array[-1]["modifiedWriteValue"] = ""
        elif re.match(r'\S+1c',access_tag):
            self.__field_array[-1]["modifiedWriteValue"] = "oneToClear"            
        elif re.match(r'\S+1s',access_tag):
              self.__field_array[-1]["modifiedWriteValue"] = "oneToSet"              
        elif re.match(r'\S+1t',access_tag):
            self.__field_array[-1]["modifiedWriteValue"] = "oneToToggle"            
        elif re.match(r'\S+0c',access_tag):
            self.__field_array[-1]["modifiedWriteValue"] = "zeroToClear"            
        elif re.match(r'\S+0s',access_tag):
              self.__field_array[-1]["modifiedWriteValue"] = "zeroToSet"              
        elif re.match(r'\S+0t',access_tag):
            self.__field_array[-1]["modifiedWriteValue"] = "zeroToToggle"            
        elif re.match(r'\S+c',access_tag):
            self.__field_array[-1]["modifiedWriteValue"] = "Clear"
        elif re.match(r'\S+s$',access_tag):
              self.__field_array[-1]["modifiedWriteValue"] = "Set"
        else:
              self.__field_array[-1]["modifiedWriteValue"] = ""


        if re.match(r'RWi|RdWi',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "RO"
        elif re.match(r'^RW|RdW$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "RW"
        elif re.match(r'^RcWi$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "RC"
        elif re.match(r'^RsWi$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "RS"
        elif re.match(r'^RcW$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WRC"
        elif re.match(r'^RsW$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WRS"
        elif re.match(r'^R(d)*Wc$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WC"
        elif re.match(r'^R(d)*Ws$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WS"
        elif re.match(r'^RcWs$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WSRC"
        elif re.match(r'^RsWc$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WCRS"
        elif re.match(r'^R(d)*W0i1c$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W1C"
        elif re.match(r'^R(d)*W0i1s$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W1S"
        elif re.match(r'^R(d)*W0i1t$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W1T"
        elif re.match(r'^R(d)*W0i0c$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W0C"
        elif re.match(r'^R(d)*W0i0s$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W0S"
        elif re.match(r'^R(d)*W0i0t$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W0T"
        elif re.match(r'^R(d)*cW0i1s$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W1SRC"
        elif re.match(r'^RsW0i1c$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W1CRS"
        elif re.match(r'^R(d)*cW0s1i$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W0SRC"
        elif re.match(r'^RsW0c1i$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W0CRS"
        elif re.match(r'^RiW$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WO"
        elif re.match(r'^RiWc$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WOC"
        elif re.match(r'^RiWs$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WOS"
        elif re.match(r'^R(d)*Wsty$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "W1"
        elif re.match(r'^RiWsty$',access_tag):
            self.__field_array[-1]["UVM_ACCESS"] = "WO1"
        else:
            assert False,"UVM_ACCESS is not matching : "+access_tag
        

        """
        TODO: Need to work on access types
        """
        if re.match(r'\S+Wi',access_tag):
           self.__field_array[-1]["Access"] = "read-only"        
        elif re.match(r'RdWsty',access_tag):
           self.__field_array[-1]["Access"] = "read-WriteOnce"
        elif re.match(r'RiWsty',access_tag):
           self.__field_array[-1]["Access"] = "writeOnce"
        elif re.match(r'Ri',access_tag):
           self.__field_array[-1]["Access"] = "write-only"   
        else:
           self.__field_array[-1]["Access"] = "read-write"

        if (access_tag == "RdWi"):
            assert self.__field_array[-1]["Access"] == "read-only", "Access tag is proper "+name+":"+access_tag+":"+self.__field_array[-1]["Access"]


        if self.__field_array[-1]["Access"] == "read-only":
            self.__field_array[-1]["sfr_doc_type"] = "RO"
        elif re.match(r'RiWsty',access_tag):
            self.__field_array[-1]["sfr_doc_type"] = "WO"
        elif self.__field_array[-1]["modifiedWriteValue"] == "oneToClear":
            self.__field_array[-1]["sfr_doc_type"] = "RW1C"
        elif self.__field_array[-1]["modifiedWriteValue"] == "oneToSet":
            self.__field_array[-1]["sfr_doc_type"] = "RW1S"
        elif self.__field_array[-1]["modifiedWriteValue"] == "oneToToggle":
            self.__field_array[-1]["sfr_doc_type"] = "RW1T"
        elif self.__field_array[-1]["modifiedWriteValue"] == "zeroToClear":
            self.__field_array[-1]["sfr_doc_type"] = "RW0C"
        elif self.__field_array[-1]["modifiedWriteValue"] == "zeroToSet":
            self.__field_array[-1]["sfr_doc_type"] = "RW0S"
        elif self.__field_array[-1]["modifiedWriteValue"] == "zeroToToggle":
            self.__field_array[-1]["sfr_doc_type"] = "RW0T"
        elif self.__field_array[-1]["modifiedWriteValue"] == "Clear":
            self.__field_array[-1]["sfr_doc_type"] = "RWC"
        elif self.__field_array[-1]["modifiedWriteValue"] == "Set":
            self.__field_array[-1]["sfr_doc_type"] = "RWS"
        else:            
            self.__field_array[-1]["sfr_doc_type"] = "RW"        
        


        if re.match(r'Rc',access_tag):
            self.__field_array[-1]["ReadAction"] = "clear"
        elif re.match(r'Rs',access_tag):
            self.__field_array[-1]["ReadAction"] = "set"
        else:
            self.__field_array[-1]["ReadAction"] = ""


        '''
        TODO:
        This below code is to find the reset value in IP-XACT.
        But this code is not yet developed completely
        '''
        m = re.match('0x([0-9a-fA-F]+)',reset_value)
        if m:
            reset =format(int(m.group(1),16),"032b")
        j=[int(d) for d in str(reset)[:]]
        reset = j
        j = 0    
        for i in range(int(lsb),int(msb)+1):
            self.__field_names[i] = name
            self.__reset_values[i] = reset[j]
            j = j + 1
        #print("name : %s, reset_value : %s, reset : %s" %(name, reset, str(self.__reset_values)))
        self.__register["reset"] = self.__reset_values

            

    def check_if_2_adj_fields_same_name(self,i):        
        while i < 32:            
            if i == 31:                
                return i
            elif self.__field_names[i] != self.__field_names[i+1] :                
                return i
            else:
                i = self.check_if_2_adj_fields_same_name(i+1)


    def convert_false_to_reserved(self):
        for i in range(len(self.__field_names)): # if (self.__field_names[i] == False):
            if (self.__field_names[i] == False):
              self.assign_reserved(i)
              



    def create_unique_names(self):                
        for i in self.__field_names:
            if i not in self.__unique_field_names and i != False:
                self.__unique_field_names.append(i)

        


    def assign_reserved(self,i):
        j = 0;
        while True:
            if "RESERVED"+str(j) not in self.__field_names:
                new_reserved_name = "RESERVED"+str(j)
                break
            j += 1
            
        if i <= 31:
            next_no_repeatation_value = self.check_if_2_adj_fields_same_name(i)            
            for k in range(i,next_no_repeatation_value+1) :
                self.__field_names[k] = new_reserved_name
            self.update_field_name(new_reserved_name,i,next_no_repeatation_value,"RdWi","0x0","","","","Reserved Field","RESERVED")            
        else:
            assert False,"Shouldn't enter else"
        return

      
    def csv_file_generation(self,file_handler):
        self.writing_into_xml_file(file_handler, "CSV,, "+self.__register["name"]+", "+str(self.__sfr_width)+", RW, "+str(self.__register["offset"])+", "+self.__register["description"]+",")
        self.writing_into_xml_file(file_handler, ",ATTR,RESET_MASK,0xffffffff")
        raw = r'"""'
        for i in self.__unique_field_names[::-1]:            
            fields = [d for d in self.__field_array if d["name"] is i]            
            field = fields[0]
            if (field['RESERVED'] != ""):
                self.writing_into_xml_file(file_handler, ",FIELD, RESERVED, "+str(field["msb"])+", "+str(field["lsb"])+", "+field["UVM_ACCESS"]+", , "+raw+field["description"]+raw+",")
                #self.writing_into_xml_file(file_handler, ",ATTR, DIRECT_READ, 1, "+ +", 0,"
                                           
        #self.writing_into_xml_file(file_handler,"\n\n\n")



    def rtl_code_header(self,file_handler):

        localtime = time.asctime( time.localtime(time.time()) )
        self.writing_into_xml_file(file,"\n\n// This is Auto Generated Code\n")
        self.writing_into_xml_file(file,"// Scripting Language    : Python")
        self.writing_into_xml_file(file,"// Developed By          : Rahul Harihara Iyer")
        self.writing_into_xml_file(file,"// Module Generated Date : "+localtime)
        self.writing_into_xml_file(file,"\n\n")
            


    def writing_c_header(self,file_handler):
        localtime = time.asctime( time.localtime(time.time()) )
        self.writing_into_xml_file(file,"\n\n// This is Auto Generated Code\n")
        self.writing_into_xml_file(file,"// Scripting Language    : Python")
        self.writing_into_xml_file(file,"// Developed By          : Rahul Harihara Iyer")
        self.writing_into_xml_file(file,"// Module Generated Date : "+localtime)
        self.writing_into_xml_file(file,"\n\n")

    def get_register_name(self,variable):
        return self.__register[variable]



    def print_fields(self,ranges,values_xls,sheet,merge_format_3):
        for k in self.__field_array:
            if (k["lsb"] == k["msb"] or int(k["lsb"]) == 15 or int(k["msb"]) == 16):
                if (int(k["lsb"]) <16):                    
                    sheet.write(values_xls[int(k["lsb"])]+str(ranges+1),k["name"], merge_format_3)
                else:
                    sheet.write(values_xls[int(k["lsb"])]+str(ranges),k["name"], merge_format_3)            
            elif (int(k["msb"]) > 15 and int(k["lsb"]) > 15):
                sheet.merge_range(values_xls[int(k["msb"])]+str(ranges)+":"+values_xls[int(k["lsb"])]+str(ranges), k["name"], merge_format_3)
            elif ((int(k["msb"]) < 16 and int(k["lsb"]) < 16)):
                sheet.merge_range(values_xls[int(k["msb"])]+str(ranges+1)+":"+values_xls[int(k["lsb"])]+str(ranges+1), k["name"], merge_format_3)
            else:
                if (int(k["msb"]) > 15):
                    assert int(k["msb"]) > 16,k["msb"]
                    sheet.merge_range(values_xls[int(k["msb"])]+str(ranges)+":"+values_xls[16]+str(ranges), k["name"], merge_format_3)
                if (int(k["lsb"]) != 15):
                    sheet.merge_range(values_xls[15]+str(ranges+1)+":"+values_xls[int(k["lsb"])]+str(ranges+1), k["name"], merge_format_3)
                
                
    def return_array_group_name(self):
        if self.__register["array"] > 1:
            return self.__register["group_name"]
        else:
            return ""
             

def create_xls_sfr_header(choose):
    if choose:

        
        values_xls ={31:'C',30:'D',29:'E',28:'F',27:'G',26:'H',25:'I',24:'J',23:'K',22:'L',21:'M',20:'N',19:'O',18:'P',17:'Q',16:'R',15:'C',14:'D',13:'E',12:'F',11:'G',10:'H',9:'I',8:'J',7:'K',6:'L',5:'M',4:'N',3:'O',2:'P',1:'Q',0:'R'}
        workbook = xlsxwriter.Workbook(module_name+".xlsx")
        sheet = workbook.add_worksheet(module_name+"_CHeader")

        border_format = workbook.add_format({    
        'border': 1,
        'align': 'center',
        'text_wrap': 1,
        'valign': 'vcenter'})
        
        merge_format_1 = workbook.add_format({
        'bold': 1,
        'border': 1,
        'align': 'center',
        'text_wrap': 1,
        'valign': 'vcenter',
        'fg_color': 'green'})
        
        merge_format_2 = workbook.add_format({
        'bold': 1,
        'border': 1,
        'align': 'center',
        'valign': 'vcenter',
        'fg_color': 'yellow'})
        
        merge_format_3 = workbook.add_format({
        'bold': 1,
        'border': 1,
        'align': 'center',
        'valign': 'vcenter',
        'fg_color': 'gray'})
        
        wrap_text = workbook.add_format({
        'bold': 0,
        'border': 1,
        'text_wrap': 1,
        'align': 'distributed',
        'valign': 'vcenter',
        'shrink': 0,
        'fg_color': 'yellow'})
        
        wrap_text_1 = workbook.add_format({
        'bold': 0,
        'border': 1,
        'text_wrap': 1,
        'italic': 1,
        'align': 'center',
        'valign': 'vcenter',
        'fg_color': 'silver'})
        
        sheet.merge_range('A1:A3', 'Addr(Hex)', merge_format_1)
        sheet.merge_range('B1:B3', 'RGSTR', merge_format_1)
        sheet.merge_range('C1:R1', 'Descriptor', merge_format_1)
        sheet.merge_range('S1:S3', 'R/W', merge_format_1)
        sheet.merge_range('T1:T3', 'Reset Value (Hex)', merge_format_1)
        sheet.merge_range('U1:U3', 'Disable', merge_format_1)
        sheet.merge_range('V1:V3', 'Size (Default: 32)', merge_format_1)
        sheet.merge_range('W1:W3', 'SFR Union', merge_format_1)
        sheet.merge_range('X1:X3', 'SFR Group Name', merge_format_1)
        sheet.merge_range('Y1:Y3', 'SFR Group', merge_format_1)
        for i in values_xls:        
            if (i <16):
                sheet.write(values_xls[i]+'3',i, merge_format_3)
            else:
                sheet.write(values_xls[i]+'2',i, merge_format_3)
                
        i = 0
        width = 0
        for j in register_object:
            ranges = 4+(i*2)
            sheet.merge_range('A'+str(ranges)+":A"+str(ranges+1), "0x"+j.get_register_name("offset"), wrap_text)
            sheet.merge_range('B'+str(ranges)+":B"+str(ranges+1), j.get_register_name("name"), wrap_text)
            sheet.merge_range('S'+str(ranges)+":S"+str(ranges+1), "", border_format)
            sheet.merge_range('T'+str(ranges)+":T"+str(ranges+1), "", border_format)
            sheet.merge_range('U'+str(ranges)+":U"+str(ranges+1), "", border_format)
            sheet.merge_range('V'+str(ranges)+":V"+str(ranges+1), "", border_format)
            sheet.merge_range('W'+str(ranges)+":W"+str(ranges+1), "", border_format)
            sheet.merge_range('X'+str(ranges)+":X"+str(ranges+1), "", border_format)
            sheet.merge_range('Y'+str(ranges)+":Y"+str(ranges+1), "", border_format)
            port  = j.get_register_name("port")
            array = j.get_register_name("s_array")
            
            '''Intent to create Border in all the fields'''
            for k in range(16):
                sheet.write(values_xls[k]+str(ranges),"", border_format)
                sheet.write(values_xls[k+16]+str(ranges+1),"", border_format)
                
            if (len(j.get_register_name("name")) > width):
                width = len(j.get_register_name("name"))
            j.print_fields(ranges,values_xls,sheet,wrap_text_1)        
            
            if (j.get_register_name("group_name") != ""):
                sheet.merge_range('X'+str(ranges)+":X"+str(ranges+1), j.get_register_name("group_name"), wrap_text_1)
                sheet.merge_range('Y'+str(ranges)+":Y"+str(ranges+1), str(j.get_register_name("port")), wrap_text_1)
                
            '''
            In case of C Header file writing, arrays are changed to Groups.
            '''
            if (j.get_register_name("e_array") != "" and j.get_register_name("next_addr") == 4):
                sheet.merge_range('W'+str(ranges)+":W"+str(ranges+1), str(j.get_register_name("e_array") - j.get_register_name("s_array") + 1) , wrap_text_1)
                logging.debug("Offset is : %s" %(j.get_register_name("offset")))
                
            i=i+1
        workbook.close()


def finding_the_title_bar(sheets):
    '''
    This function finds the exact row and column number
    '''
    row_column_addr = [-1,-1]
    title_bar = ['Port', 'Group_Name', 'DRepeat','Reg. Name','Offset','FIELD NAME','Range','_customType_','_hwAccess_','_sysrdl_precedence_','Reset Mask','Testable','Constraint','Byte Access Support','Bus','Coverage','Coverage Bins','HDL Path','Soft Reset Mask','Description']
    for i in range(worksheet.nrows):
        if title_bar[0] in worksheet.row_values(i) :
            reg_port     = worksheet.row_values(i).index(title_bar[0])
            row_column_addr = [i,reg_port]
            return(row_column_addr)
    assert False,"The Table Column Names is not matching"
    return(-1)



if sys.argv.count("-help") == 1:
    print("\nFormat of input is sfr_file.exe -xls <xls> -sheet <sheet_name> [-m <module_name>]")
    print("\nProcedure to run the exe file.")
    print("==============================\n")
    print("1. Copy the sfr xls whose format is in EDM (Controller Development Team > Manual_Share_SRIB > RGX > 2_Design > IP_EXACT_SFR.xlsx)")
    print("2. Open Microsoft Outlook and click New E-Mail")
    print("3. Attach the xls file in the mail.")
    print("4. Right click on the attached file and save as a different name in your PC. (Remember step 2 to 4 is a work around the NASCA)")
    print("5. Open command prompt (Windows -> RUN -> cmd)")
    print("6. Navigate to the path where the exe file is saved.")
    print("7. Type the command (sfr_file.exe -xls <File saved in Step-4> -sheet <> -m <> -data_width <> -strobe)")
    print("In Step-7, -m is the module name of the verilog file to be generated")
    sys.exit(1)
    
if sys.argv.count("-xls") == 1:
    xls_index = sys.argv.index("-xls")
    workbook_name = sys.argv[xls_index + 1]    
    assert os.path.exists(workbook_name), "Workbook path does not exists"
    assert os.path.isfile(workbook_name), "Workbook should be a file not a folder"    
else:
    raise ValueError("Format of input is sfr_code.py -xls <xls> -sheet <sheet_name> -m <> -data_width <>")

if sys.argv.count("-sheet") == 1:
    sheet_name_index = sys.argv.index("-sheet")
    worksheet_name = sys.argv[sheet_name_index + 1]
else:
    raise ValueError("Format of input is sfr_code.py -xls <xls> -sheet <sheet_name> -m <> -data_width <>")

if sys.argv.count("-m") > 0:
    module_name = sys.argv[sys.argv.index("-m")+1].lower()
else:
    module_name = worksheet_name.lower()

if sys.argv.count("-data_width") > 0:
    data_width = int(sys.argv[sys.argv.index("-data_width")+1])
else:
    data_width = 32

if sys.argv.count("-addr_width") > 0:
    addr_width = int(sys.argv[sys.argv.index("-addr_width")+1])
else:
    addr_width = 32

if sys.argv.count("-strobe") > 0:
    strobe = 1
else:
    strobe = 0

	
date = str(datetime.date.today())
new_folder_name = "sfr_files_" + date

python_file_location = os.getcwd()

workbook_name = os.path.abspath(workbook_name)
os.chdir(os.path.dirname(workbook_name))


reg_name_num = 1


register_start = False
create_reserved_fields = False

workbook = xlrd.open_workbook(workbook_name)
worksheet = workbook.sheet_by_name(worksheet_name)
row_column_addr = finding_the_title_bar(worksheet)
row_number = row_column_addr[0]
reg_port = row_column_addr[1]
assert (row_number >= 0 and reg_name_num >= 0),"It cannot be less than 0"

#reg_array = reg_port + 1
reg_group_name = reg_port + 1
reg_drepeat   = reg_group_name + 1
reg_name_num = reg_drepeat + 1
offset_num = reg_name_num + 1
field_name_num = offset_num + 1

range_num  = offset_num + 2
access     = range_num + 1
reset      = access + 1
sysrdl     = reset + 3
description_num = sysrdl + 10
reg_num    = -1;

if not os.path.isdir(new_folder_name):    
    os.mkdir(new_folder_name)    
    
os.chdir(new_folder_name)


try:
    f_handler = open(module_name+".xml","w")
except:
    raise ValueError("%s.xml does not exists" %(module_name))

'''
Loop till end of the xls file
'''
register_object = []

for i in range(row_number+1,worksheet.nrows):
    if (worksheet.cell_value(i,reg_name_num) != '') :
        if (register_start == True):
            if create_reserved_fields == True:
                register_object[reg_num].convert_false_to_reserved()
            register_object[reg_num].create_unique_names()                
        reg_name = str(worksheet.cell_value(i,reg_name_num)).strip(' ')
        assert not re.match(r'\s+',reg_name), "Register name has space : Register name : %s" %(reg_name)
        offset = str(worksheet.cell_value(i,offset_num)).strip(' ')
        assert not re.match(r'\s+',offset), "Offset shouldn't have spaces : Register name : %s" %(reg_name)
        array = worksheet.cell_value(i,field_name_num)
        
        if (str(worksheet.cell_value(i,reg_port)).strip(' ')) == "":
            port = 1
        else:
            port = int(worksheet.cell_value(i,reg_port))

        group_name = str(worksheet.cell_value(i,reg_group_name)).strip(' ')        
        drepeat    = str(worksheet.cell_value(i,reg_drepeat)).strip(' ')
            
        if (reg_name == ""):
            continue
        assert not reg_name in Register.register_name,"Register Name exists : %s" %(reg_name)
        assert not reg_name in Register.register_address,"Register Offset Address exists : %s : %s" %(reg_name,register_address)
        assert worksheet.cell_value(i,field_name_num) == '' or re.match(r'\d+\.{2}\d+',worksheet.cell_value(i,field_name_num)),"Field Name should be empty in register space. Register name : %s" % (reg_name)        
        desc   = str(worksheet.cell_value(i,description_num))
        register_object.append(Register(f_handler,32))
        register_object[reg_num].update_register_name(reg_name, offset, array, desc, port, group_name, drepeat)
        Register.register_name.append(reg_name)
        Register.register_name.append(offset)        
        register_start = True        
        continue

    
    if (register_start == True and worksheet.cell_value(i,field_name_num) != '') :
        field_name = str(worksheet.cell_value(i,field_name_num)).strip(' ')
        assert not re.match(r'\s+',field_name), "field_name shouldn't have spaces : Register name : %s, Field Name : %s" %(reg_name, field_name)
        access_tag = str(worksheet.cell_value(i,access))
        reset_value = str(worksheet.cell_value(i,reset))
        sysrdl_precedence = str(worksheet.cell_value(i,sysrdl))
        desc   = str(worksheet.cell_value(i,description_num))
                
        m = re.match('\[(.*):(.*)\]',worksheet.cell_value(i,range_num))
        if m:
            msb = m.group(1)
            lsb = m.group(2)            
        n = re.match(r'\[(.*)\]',worksheet.cell_value(i,range_num))
        if n and not m:
            bit_offset = n.group(1)
            lsb = n.group(1)
            msb = n.group(1)        

        reset_value = str(worksheet.cell_value(i,reset))
        register_object[reg_num].update_field_name(field_name,lsb,msb,access_tag,reset_value,sysrdl_precedence,"","",desc,"")



if (register_start == True):
    if create_reserved_fields == True:
        register_object[reg_num].convert_false_to_reserved()
    register_object[reg_num].create_unique_names()
    
for i in register_object:
    i.printing()
    
register_start == False



create_xls_sfr_header(True)


    
#doc=docx.Document()
for i in register_object:
    i.print_register_names(f_handler)

'''
    Below are for later use to generate CSV file ans Doc
    try:
        csv_handler = open("csv_file.csv","w")
    except:
        print("File cannot be opened")        
    i.csv_file_generation(csv_handler)
    i.writing_sfr_doc(doc,offset_address)
    csv_handler.close()
'''

f_handler.close()
sfr_code.sfr_verilog_code(module_name+".xml",module_name, data_width, addr_width, strobe)
print("Files created in %s folder are as follows" %(os.getcwd()))
for i in os.listdir(os.getcwd()):
    print(i)
